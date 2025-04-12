import os
import uuid
import logging
from werkzeug.utils import secure_filename
from flask import current_app
from ..models import db, Book

# Import file format handlers
import ebooklib
from ebooklib import epub
from PyPDF2 import PdfReader
import docx
from bs4 import BeautifulSoup

# Импорт модулей для обработки файлов...

logger = logging.getLogger(__name__)


class FileService:
    """Service for handling file uploads and processing"""

    def __init__(self):
        # Не обращаемся к current_app здесь
        pass

    def allowed_file(self, filename):
        """Check if file has an allowed extension"""
        from flask import current_app
        allowed_extensions = current_app.config.get('ALLOWED_EXTENSIONS',
                                                    {'txt', 'pdf', 'epub', 'mobi', 'doc', 'docx', 'rtf'})
        return '.' in filename and filename.rsplit('.', 1)[1].lower() in allowed_extensions

    def save_file(self, file, user_id):
        """Save an uploaded file and create a Book record"""
        if not file or not self.allowed_file(file.filename):
            return None

        try:
            from flask import current_app
            import uuid
            import os

            # Получаем путь к папке загрузок из конфигурации
            upload_folder = current_app.config.get('UPLOAD_FOLDER', 'static/uploads')

            # Generate a unique filename
            original_filename = secure_filename(file.filename)
            file_extension = original_filename.rsplit('.', 1)[1].lower()
            unique_filename = f"{str(uuid.uuid4())}.{file_extension}"

            # Create user's directory if it doesn't exist
            user_upload_dir = os.path.join(upload_folder, str(user_id))
            os.makedirs(user_upload_dir, exist_ok=True)

            # Save the file
            file_path = os.path.join(user_upload_dir, unique_filename)
            file.save(file_path)

            # Extract metadata
            metadata = self.extract_metadata(file_path, file_extension)

            # Убедитесь, что title не является None
            title = metadata.get('title')
            if not title:
                title = original_filename  # Используйте имя файла если нет заголовка

            # Create Book record
            book = Book(
                title=title,
                author=metadata.get('author') or 'Unknown',
                language=metadata.get('language') or current_app.config.get('DEFAULT_LANGUAGE', 'en'),
                file_path=file_path,
                original_filename=original_filename,
                file_type=file_extension,
                word_count=metadata.get('word_count', 0),
                user_id=user_id
            )

            db.session.add(book)
            db.session.commit()

            return book

        except Exception as e:
            current_app.logger.error(f"Error saving file: {e}")
            db.session.rollback()
            # Remove file if it was saved
            if 'file_path' in locals() and os.path.exists(file_path):
                os.remove(file_path)
            return None

    def _extract_docx_metadata(self, file_path):
        """Extract metadata from DOCX file"""
        metadata = {
            'title': None,
            'author': None,
            'language': None,
            'word_count': 0
        }

        try:
            doc = docx.Document(file_path)
            docx_properties = doc.core_properties

            metadata['title'] = docx_properties.title or os.path.basename(file_path).rsplit('.', 1)[0]
            metadata['author'] = docx_properties.author or 'Unknown'

            # Count words
            word_count = 0
            for para in doc.paragraphs:
                word_count += len(para.text.split())

            metadata['word_count'] = word_count
            return metadata

        except Exception as e:
            logger.error(f"Error extracting DOCX metadata: {e}")
            # Вернем имя файла в качестве заголовка в случае ошибки
            metadata['title'] = os.path.basename(file_path).rsplit('.', 1)[0]
            return metadata

    def _extract_txt_metadata(self, file_path):
        """Extract metadata from TXT file"""
        metadata = {
            'title': None,
            'author': None,
            'language': None,
            'word_count': 0
        }

        try:
            # For text files, use the filename as title
            metadata['title'] = os.path.basename(file_path).rsplit('.', 1)[0]

            # Count words
            with open(file_path, 'r', encoding='utf-8', errors='replace') as file:
                text = file.read()
                metadata['word_count'] = len(text.split())

            return metadata

        except Exception as e:
            logger.error(f"Error extracting TXT metadata: {e}")
            return metadata

    def get_text_content(self, book):
        """
        Extract text content from a book file

        Args:
            book: Book object with file path information

        Returns:
            Extracted text content
        """
        file_path = book.file_path
        file_extension = book.file_type

        try:
            if file_extension == 'pdf':
                return self._extract_pdf_text(file_path)
            elif file_extension in ['epub', 'mobi']:
                return self._extract_epub_text(file_path)
            elif file_extension in ['doc', 'docx']:
                return self._extract_docx_text(file_path)
            elif file_extension == 'txt':
                return self._extract_txt_text(file_path)
            else:
                return "Unsupported file format."

        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {e}")
            return f"Error extracting text: {str(e)}"

    def _extract_pdf_text(self, file_path):
        """Extract text from PDF file"""
        text = ""
        try:
            with open(file_path, 'rb') as file:
                reader = PdfReader(file)
                for page in reader.pages:
                    text += page.extract_text() + "\n\n"
            return text
        except Exception as e:
            logger.error(f"Error extracting PDF text: {e}")
            return f"Error extracting text: {str(e)}"

    def _extract_epub_text(self, file_path):
        """Extract text from EPUB file"""
        text = ""
        try:
            book = epub.read_epub(file_path)
            for item in book.get_items_of_type(ebooklib.ITEM_DOCUMENT):
                content = item.get_content().decode('utf-8')
                soup = BeautifulSoup(content, 'html.parser')
                text += soup.get_text() + "\n\n"
            return text
        except Exception as e:
            logger.error(f"Error extracting EPUB text: {e}")
            return f"Error extracting text: {str(e)}"

    def _extract_docx_text(self, file_path):
        """Extract text from DOCX file"""
        text = ""
        try:
            doc = docx.Document(file_path)
            for para in doc.paragraphs:
                text += para.text + "\n"
            return text
        except Exception as e:
            logger.error(f"Error extracting DOCX text: {e}")
            return f"Error extracting text: {str(e)}"

    def _extract_txt_text(self, file_path):
        """Extract text from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='replace') as file:
                return file.read()
        except Exception as e:
            logger.error(f"Error extracting TXT text: {e}")
            return f"Error extracting text: {str(e)}"

    def format_text_for_reader(self, text):
        """
        Format text for the reader view

        Args:
            text: The raw text content

        Returns:
            HTML-formatted text with clickable words
        """
        # Split text into words and wrap each in a clickable span
        words = text.split()
        wrapped_words = []

        for word in words:
            # Skip empty strings and punctuation-only strings
            if not word or not any(c.isalnum() for c in word):
                wrapped_words.append(word)
                continue

            # Strip punctuation for the word lookup but keep it in the display
            clean_word = word.strip('.,;:!?()[]{}""\'')
            if not clean_word:
                wrapped_words.append(word)
                continue

            # Create a clickable span
            wrapped_word = f'<span class="word" data-word="{clean_word}">{word}</span>'
            wrapped_words.append(wrapped_word)

        # Join words back together and replace newlines with <br> tags
        formatted_text = ' '.join(wrapped_words)
        formatted_text = formatted_text.replace('\n', '<br>')

        return formatted_text

    def extract_metadata(self, file_path, file_extension):
        """
        Extract metadata from a file based on its type

        Args:
            file_path: Path to the file
            file_extension: Extension of the file

        Returns:
            Dictionary with metadata
        """
        metadata = {
            'title': None,
            'author': None,
            'language': None,
            'word_count': 0
        }

        try:
            if file_extension == 'pdf':
                return self._extract_pdf_metadata(file_path)
            elif file_extension in ['epub', 'mobi']:
                return self._extract_epub_metadata(file_path)
            elif file_extension in ['doc', 'docx']:
                return self._extract_docx_metadata(file_path)
            elif file_extension == 'txt':
                return self._extract_txt_metadata(file_path)
            else:
                return metadata

        except Exception as e:
            logger.error(f"Error extracting metadata from {file_path}: {e}")
            return metadata

    def _extract_pdf_metadata(self, file_path):
        """Extract metadata from PDF file"""
        metadata = {
            'title': None,
            'author': None,
            'language': None,
            'word_count': 0
        }

        try:
            with open(file_path, 'rb') as file:
                reader = PdfReader(file)
                info = reader.metadata

                if info:
                    metadata['title'] = info.get('/Title', None)
                    metadata['author'] = info.get('/Author', None)

                # Extract text and count words
                text = ""
                for page in reader.pages:
                    text += page.extract_text() + "\n"

                metadata['word_count'] = len(text.split())
                return metadata

        except Exception as e:
            logger.error(f"Error extracting PDF metadata: {e}")
            return metadata

    def _extract_epub_metadata(self, file_path):
        """Extract metadata from EPUB file"""
        metadata = {
            'title': None,
            'author': None,
            'language': None,
            'word_count': 0
        }

        try:
            book = epub.read_epub(file_path)

            metadata['title'] = (book.get_metadata('DC', 'title')[0][0]
                                 if book.get_metadata('DC', 'title') else None)
            metadata['author'] = (book.get_metadata('DC', 'creator')[0][0]
                                  if book.get_metadata('DC', 'creator') else None)
            metadata['language'] = (book.get_metadata('DC', 'language')[0][0]
                                    if book.get_metadata('DC', 'language') else None)

            # Count words in the book
            word_count = 0
            for item in book.get_items_of_type(ebooklib.ITEM_DOCUMENT):
                content = item.get_content().decode('utf-8')
                soup = BeautifulSoup(content, 'html.parser')
                text = soup.get_text()
                word_count += len(text.split())

            metadata['word_count'] = word_count
            return metadata

        except Exception as e:
            logger.error(f"Error extracting EPUB metadata: {e}")
            return metadata
